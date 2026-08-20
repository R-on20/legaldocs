"""
Word Document to HTML Converter and Editor
"""
import io
import tempfile
from typing import Optional, Tuple
from docx import Document as DocxDocument
from docx.shared import Inches
import mammoth
from bs4 import BeautifulSoup
import re
from django.core.files.base import ContentFile


class DocxToHtmlConverter:
    """Convert DOCX files to HTML and back"""
    
    def docx_to_html(self, file_path_or_content) -> str:
        """
        Convert DOCX file to HTML using mammoth library
        
        Args:
            file_path_or_content: File path string or file content bytes
            
        Returns:
            HTML string
        """
        try:
            if isinstance(file_path_or_content, str):
                # File path
                with open(file_path_or_content, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
            else:
                # File content (bytes)
                result = mammoth.convert_to_html(io.BytesIO(file_path_or_content))
            
            html_content = result.value
            
            # Check if we got meaningful content
            if not html_content or len(html_content.strip()) < 10:
                # Fallback to basic extraction
                return self._fallback_conversion(file_path_or_content)
            
            # Clean up the HTML for better CKEditor compatibility
            html_content = self._clean_html_for_editor(html_content)
            
            return html_content
            
        except Exception as e:
            print(f"Error converting DOCX to HTML: {e}")
            # Try fallback method
            return self._fallback_conversion(file_path_or_content)
    
    def _fallback_conversion(self, file_path_or_content):
        """Fallback conversion using python-docx"""
        try:
            from docx import Document as DocxDocument
            
            if isinstance(file_path_or_content, str):
                doc = DocxDocument(file_path_or_content)
            else:
                doc = DocxDocument(io.BytesIO(file_path_or_content))
            
            html_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Check for heading styles
                    style_name = paragraph.style.name.lower()
                    if 'heading' in style_name:
                        level = 1  # Default heading level
                        if 'heading 1' in style_name:
                            level = 1
                        elif 'heading 2' in style_name:
                            level = 2
                        elif 'heading 3' in style_name:
                            level = 3
                        elif 'heading 4' in style_name:
                            level = 4
                        elif 'heading 5' in style_name:
                            level = 5
                        elif 'heading 6' in style_name:
                            level = 6
                        html_parts.append(f'<h{level}>{text}</h{level}>')
                    else:
                        # Format inline styling
                        formatted_text = self._format_runs(paragraph)
                        html_parts.append(f'<p>{formatted_text}</p>')
                else:
                    html_parts.append('<p><br/></p>')  # Empty paragraph
            
            return '\n'.join(html_parts) if html_parts else '<p>No content found in document</p>'
            
        except Exception as e:
            print(f"Fallback conversion failed: {e}")
            return f'<p>Error converting document: {str(e)}</p>'
    
    def _format_runs(self, paragraph):
        """Format runs within a paragraph to preserve styling"""
        formatted_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
                
            if run.bold:
                text = f'<strong>{text}</strong>'
            if run.italic:
                text = f'<em>{text}</em>'
            if run.underline:
                text = f'<u>{text}</u>'
                
            formatted_parts.append(text)
        
        return ''.join(formatted_parts)
    
    def html_to_docx(self, html_content: str, title: str = "Document") -> bytes:
        """
        Convert HTML content to DOCX format
        
        Args:
            html_content: HTML string
            title: Document title
            
        Returns:
            DOCX file content as bytes
        """
        try:
            # Create a new Document
            doc = DocxDocument()
            
            # Don't add title as heading to avoid duplication
            # The title will be in the document metadata instead
            
            # Parse HTML content
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Convert HTML elements to DOCX
            self._html_to_docx_elements(soup, doc)
            
            # If no content was added, add a basic paragraph
            if len(doc.paragraphs) == 0:
                doc.add_paragraph("Document content")
            
            # Save to BytesIO
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except Exception as e:
            print(f"Error converting HTML to DOCX: {e}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
            # Return a simple document with the content as text
            try:
                doc = DocxDocument()
                # Strip HTML tags and add as plain text
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')
                plain_text = soup.get_text()
                if plain_text.strip():
                    doc.add_paragraph(plain_text)
                else:
                    doc.add_paragraph("Error: No content to convert")
                
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                return doc_buffer.getvalue()
            except Exception as fallback_error:
                print(f"Fallback conversion also failed: {fallback_error}")
                return b''  # Return empty bytes if everything fails
    
    def _clean_html_for_editor(self, html_content: str) -> str:
        """Clean HTML content for better CKEditor compatibility and rendering"""
        # Parse with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove empty paragraphs but keep paragraphs with just whitespace for spacing
        for p in soup.find_all('p'):
            text_content = p.get_text()
            # Only remove completely empty paragraphs (no text, no &nbsp;, no images, etc.)
            if not text_content and not p.find_all(['img', 'br']) and '&nbsp;' not in str(p):
                p.decompose()
        
        # Improve table spacing and formatting
        for table in soup.find_all('table'):
            # Add proper table classes for better styling
            table['class'] = table.get('class', []) + ['table', 'table-bordered']
            table['style'] = 'margin: 1em 0; border-collapse: collapse; width: 100%;'
            
            # Ensure proper cell spacing
            for cell in table.find_all(['td', 'th']):
                if not cell.get('style'):
                    cell['style'] = 'padding: 8px; border: 1px solid #ddd;'
                else:
                    # Add padding if not present
                    if 'padding' not in cell['style']:
                        cell['style'] += '; padding: 8px;'
        
        # Fix spacing around tables - add paragraphs before/after if needed
        for table in soup.find_all('table'):
            # Check if table has proper spacing before
            prev_sibling = table.previous_sibling
            while prev_sibling and str(prev_sibling).strip() == '':
                prev_sibling = prev_sibling.previous_sibling
            
            if prev_sibling and hasattr(prev_sibling, 'name') and prev_sibling.name not in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Insert a paragraph for spacing
                spacing_p = soup.new_tag('p')
                spacing_p.string = '\u00A0'  # Non-breaking space
                table.insert_before(spacing_p)
            
            # Check if table has proper spacing after
            next_sibling = table.next_sibling
            while next_sibling and str(next_sibling).strip() == '':
                next_sibling = next_sibling.next_sibling
                
            if next_sibling and hasattr(next_sibling, 'name') and next_sibling.name not in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Insert a paragraph for spacing
                spacing_p = soup.new_tag('p')
                spacing_p.string = '\u00A0'  # Non-breaking space
                table.insert_after(spacing_p)
        
        # Convert some elements for better compatibility but preserve formatting
        # Keep <strong> and <em> as they are more semantic
        
        # Clean up excessive whitespace but preserve intentional spacing
        html_str = str(soup)
        
        # Remove excessive empty lines but keep some spacing
        html_str = re.sub(r'\n\s*\n\s*\n+', '\n\n', html_str)
        
        # Ensure proper paragraph structure if needed
        soup = BeautifulSoup(html_str, 'html.parser')
        if not soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'div', 'table']):
            # If no block elements, wrap content in paragraphs
            content = soup.get_text()
            lines = content.split('\n')
            new_soup = BeautifulSoup('', 'html.parser')
            for line in lines:
                if line.strip():
                    p = new_soup.new_tag('p')
                    p.string = line.strip()
                    new_soup.append(p)
            soup = new_soup
        
        return str(soup)
    
    def _html_to_docx_elements(self, soup: BeautifulSoup, doc: DocxDocument):
        """Convert HTML elements to DOCX elements"""
        
        def process_element(element, parent_paragraph=None):
            if not hasattr(element, 'name') or element.name is None:
                # Text node
                text = str(element).strip()
                if text and parent_paragraph:
                    parent_paragraph.add_run(text)
                elif text:
                    # Create a new paragraph for standalone text
                    p = doc.add_paragraph()
                    p.add_run(text)
                return
                
            if element.name == 'p':
                p = doc.add_paragraph()
                for child in element.children:
                    if hasattr(child, 'name') and child.name:
                        process_inline_element(child, p)
                    else:
                        # Text node
                        text = str(child).strip()
                        if text:
                            p.add_run(text)
            
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                text = element.get_text().strip()
                if text:
                    doc.add_heading(text, level)
            
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Bullet')
            
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Number')
            
            elif element.name == 'table':
                # Basic table support
                rows = element.find_all('tr')
                if rows:
                    cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                    if cols > 0:
                        table = doc.add_table(rows=len(rows), cols=cols)
                        
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if j < cols and i < len(table.rows):
                                    table.cell(i, j).text = cell.get_text().strip()
            
            elif element.name == 'br':
                if parent_paragraph:
                    parent_paragraph.add_run().add_break()
                else:
                    # Add empty paragraph for line break
                    doc.add_paragraph()
            
            elif element.name in ['div', 'span']:
                # Process children for container elements
                for child in element.children:
                    process_element(child, parent_paragraph)
            
            else:
                # For other elements, just process their text content
                text = element.get_text().strip()
                if text:
                    if parent_paragraph:
                        parent_paragraph.add_run(text)
                    else:
                        doc.add_paragraph(text)
        
        def process_inline_element(element, paragraph):
            """Process inline elements within a paragraph"""
            if not hasattr(element, 'name') or element.name is None:
                text = str(element).strip()
                if text:
                    paragraph.add_run(text)
                return
                
            text = element.get_text()
            if not text:
                return
                
            if element.name in ['b', 'strong']:
                run = paragraph.add_run(text)
                run.bold = True
            elif element.name in ['i', 'em']:
                run = paragraph.add_run(text)
                run.italic = True
            elif element.name == 'u':
                run = paragraph.add_run(text)
                run.underline = True
            elif element.name == 'br':
                paragraph.add_run().add_break()
            else:
                paragraph.add_run(text)
        
        # Process all top-level elements
        for element in soup.children:
            if hasattr(element, 'name') and element.name:
                process_element(element)
            else:
                # Handle text nodes at the top level
                text = str(element).strip()
                if text:
                    doc.add_paragraph(text)


def convert_docx_to_html(file_content: bytes) -> str:
    """
    Utility function to convert DOCX content to HTML
    
    Args:
        file_content: DOCX file content as bytes
        
    Returns:
        HTML string
    """
    converter = DocxToHtmlConverter()
    return converter.docx_to_html(file_content)


def convert_html_to_docx(html_content: str, title: str = "Document") -> bytes:
    """
    Utility function to convert HTML to DOCX
    
    Args:
        html_content: HTML string
        title: Document title
        
    Returns:
        DOCX file content as bytes
    """
    converter = DocxToHtmlConverter()
    return converter.html_to_docx(html_content, title)
