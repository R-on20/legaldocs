from io import BytesIO
import re
from django.shortcuts import render
from django.shortcuts import render
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView
from main.models import Document, DocumentTemplate, DocumentVersion, Tag, AuditLog
from django.views.generic import ListView, CreateView, UpdateView, DetailView, DeleteView
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from django.urls import reverse, reverse_lazy
import uuid
import tempfile
import logging
from django.db import transaction
from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect
from django.views.generic import View
from django.views.generic import TemplateView
from django.http import Http404, JsonResponse
from django.db.models import Q
from urllib.parse import unquote
from django.utils.translation import gettext_lazy as _
from django.contrib.postgres.search import SearchQuery, SearchRank, SearchVector
from django.conf import settings
from .forms import DocumentUploadForm, DocumentVersionForm, DocumentUpdateForm, DocumentSearchForm, AISearchForm, DocumentCreationForm, DocumentWordEditForm
from .utils import get_search_queryset
from .ai_search import AISearchProcessor, apply_ai_filters_to_queryset
from docx import Document as DocxDocument
from django.core.files.base import ContentFile

logger = logging.getLogger(__name__)

# Helper function for role-based permissions
def can_manage_documents(user):
    """Check if user can create/upload documents"""
    from main.models import User
    return user.is_authenticated and user.role in [User.Role.LAWYER, User.Role.ADMIN, User.Role.PARALEGAL, User.Role.SENIOR_PARALEGAL]


class DocumentProcessingStatusView(LoginRequiredMixin, View):
    """Check the processing status of a document"""
    
    def get(self, request, pk):
        document = get_object_or_404(Document, pk=pk)
        
        # Check if user has permission to view this document
        if (document.confidential and 
            not request.user.has_perm('documents.can_view_confidential')):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        status_info = {
            'document_id': document.pk,
            'title': document.title,
            'status': document.status,
            'has_content': bool(document.content),
            'has_html_content': bool(document.html_content),
            'has_transcript': bool(document.audio_transcript),
            'file_size': document.file.size if document.file else 0,
            'is_processing': document.status == Document.Status.DRAFT,
            'processing_complete': document.status != Document.Status.DRAFT,
        }
        
        return JsonResponse(status_info)


class DocumentDownloadView(LoginRequiredMixin, View):
    """Handle document downloads with audit logging"""
    
    def get(self, request, pk):
        document = get_object_or_404(Document, pk=pk)
        
        # Check if user has permission to view this document
        if (document.confidential and 
            not request.user.has_perm('documents.can_view_confidential')):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        # Log the download
        try:
            AuditLog.log_action(
                user=request.user,
                document=document,
                action='DOWNLOAD',
                request=request,
                file_name=document.file.name,
                file_size=document.file.size
            )
        except Exception as e:
            logger.warning(f"Failed to log download: {e}")
        
        # Redirect to actual file URL
        from django.http import HttpResponseRedirect
        return HttpResponseRedirect(document.file.url)

def home(request):
    """
    Home view that redirects authenticated users to dashboard
    and shows landing page for anonymous users
    """
    if request.user.is_authenticated:
        from django.shortcuts import redirect
        return redirect('main:dashboard')
    return render(request, 'index.html')


class DocumentListView(LoginRequiredMixin, ListView):
    model = Document
    template_name = 'documents/document_list.html'
    context_object_name = 'documents'
    paginate_by = 20
    search_form_class = DocumentSearchForm

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.ai_processor = AISearchProcessor()

    def get_queryset(self):
        qs = super().get_queryset()
        
        # Handle archived filtering from URL parameters
        is_archived = self.request.GET.get('is_archived')
        if is_archived == 'false':
            qs = qs.filter(is_archived=False)
        elif is_archived == 'true':
            qs = qs.filter(is_archived=True)
        else:
            # Default to non-archived if no specific filter
            qs = qs.filter(is_archived=False)
        
        if not self.request.user.has_perm('documents.can_view_confidential'):
            qs = qs.filter(confidential=False)
        
        # Handle both AI search and regular search
        ai_query = self.request.GET.get('ai_query')
        search_query = self.request.GET.get('q')
        
        # Initialize search-related attributes
        self.search_method = None
        self.ai_search_result = None
        self.search_query = None
        self.search_count = 0
        self.total_count = qs.count()
        
        if ai_query:
            # AI-powered search
            logger.info(f"AI search query received: '{ai_query}'")
            original_count = qs.count()
            logger.info(f"Total documents before AI search: {original_count}")
            
            try:
                ai_result = self.ai_processor.process_query(ai_query)
                
                if ai_result['success']:
                    logger.info(f"AI search processing successful using method: {ai_result['method']}")
                    
                    # Apply AI-extracted filters
                    qs = apply_ai_filters_to_queryset(qs, ai_result)
                    
                    # Store AI results for template context
                    self.ai_search_result = ai_result
                    self.search_method = 'ai'
                    self.search_query = ai_query
                else:
                    # Fallback to regular search with AI query
                    logger.info("AI search failed, falling back to regular search")
                    qs = get_search_queryset(qs, ai_query)
                    self.search_method = 'fallback'
                    self.search_query = ai_query
                    
            except Exception as e:
                logger.error(f"AI search error: {e}")
                # Fallback to regular search
                qs = get_search_queryset(qs, ai_query)
                self.search_method = 'fallback'
                self.search_query = ai_query
            
            search_count = qs.count()
            logger.info(f"Documents found after AI search: {search_count}")
            self.search_count = search_count
            
        elif search_query:
            # Regular keyword search
            logger.info(f"Regular search query received: '{search_query}'")
            original_count = qs.count()
            logger.info(f"Total documents before search: {original_count}")
            
            qs = get_search_queryset(qs, search_query)
            self.search_method = 'keyword'
            self.search_query = search_query
            
            search_count = qs.count()
            logger.info(f"Documents found after search: {search_count}")
            self.search_count = search_count
        
        # Additional filtering (existing form-based filters)
        document_type = self.request.GET.get('document_type')
        if document_type:
            qs = qs.filter(document_type=document_type)
            
        status = self.request.GET.get('status')
        if status:
            qs = qs.filter(status=status)
            
        uploaded_after = self.request.GET.get('uploaded_after')
        if uploaded_after:
            qs = qs.filter(uploaded_at__gte=uploaded_after)
            
        return qs

    def get(self, request, *args, **kwargs):
        """Override get method to handle AJAX requests and smart search auto-redirect"""
        
        # Handle AJAX requests for search updates
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            response = super().get(request, *args, **kwargs)
            return response
        
        ai_query = request.GET.get('ai_query')
        search_query = request.GET.get('q')
        
        # Check for single document match on either search type
        if ai_query or search_query:
            # Get the filtered queryset
            queryset = self.get_queryset()
            
            # If exactly one document matches, redirect to its detail view with search query
            if queryset.count() == 1:
                document = queryset.first()
                from urllib.parse import urlencode
                query_param = ai_query if ai_query else search_query
                param_name = 'ai_query' if ai_query else 'q'
                return redirect(f"{reverse('documents:document_detail', kwargs={'pk': document.pk})}?{urlencode({param_name: query_param})}")
        
        return super().get(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['search_form'] = self.search_form_class(self.request.GET)
        context['ai_search_form'] = AISearchForm(self.request.GET)
        
        # Add search information
        ai_query = self.request.GET.get('ai_query')
        search_query = self.request.GET.get('q')
        
        if ai_query or search_query:
            query = ai_query if ai_query else search_query
            context['search_query'] = query
            context['search_count'] = getattr(self, 'search_count', 0)
            context['total_count'] = getattr(self, 'total_count', 0)
            context['has_search_results'] = context['search_count'] > 0
            context['search_method'] = getattr(self, 'search_method', 'unknown')
            
            # Determine search type for display
            if ai_query:
                context['is_ai_search'] = True
                context['search_type_display'] = 'AI Enhanced Search'
            else:
                context['is_ai_search'] = False
                context['search_type_display'] = 'Keyword Search'
            
            # Add AI search results if available
            ai_result = getattr(self, 'ai_search_result', None)
            if ai_result:
                context['ai_search_result'] = ai_result
                context['extracted_filters'] = ai_result.get('filters', {})
                context['extracted_keywords'] = ai_result.get('keywords', [])
                
                # Add method indicator for display
                if ai_result['method'] == 'ai':
                    context['ai_processing_status'] = 'success'
                    context['ai_processing_message'] = 'AI Processing: ✅ Working!'
                else:
                    context['ai_processing_status'] = 'fallback'
                    context['ai_processing_message'] = 'AI Processing: ⚠️ Fallback to keyword search'
            
            # Debug information for development
            if settings.DEBUG:
                from .utils import extract_search_terms
                context['search_terms'] = extract_search_terms(query)
        
        return context

class DocumentCreateView(LoginRequiredMixin, View):
    template_name = 'documents/document_create.html'

    def dispatch(self, request, *args, **kwargs):
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to create documents.")
            return redirect('documents:document_list')
        return super().dispatch(request, *args, **kwargs)

    def get(self, request):
        form = DocumentCreationForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request):
        form = DocumentCreationForm(request.POST)
        if form.is_valid():
            import tempfile
            import os
            from django.core.files.base import ContentFile
            
            try:
                # Create the actual document file from CKEditor content
                html_content = form.cleaned_data['content']
                doc_title = form.cleaned_data['name']
                
                # Convert HTML to DOCX using our converter
                from documents.docx_converter import convert_html_to_docx
                docx_content = convert_html_to_docx(html_content, doc_title)
                
                if not docx_content:
                    messages.error(request, _("Failed to create document file. Please try again."))
                    return render(request, self.template_name, {'form': form})
                
                # Create the Document model instance
                document = Document(
                    title=doc_title,
                    document_type=Document.DocType.OTHER,
                    uploaded_by=request.user,
                    editable=True,
                    content=form.cleaned_data['content'],  # Store original HTML as content
                    html_content=html_content,  # Store HTML for editing
                    status=Document.Status.DRAFT
                )
                
                # Save the DOCX file
                filename = f"{doc_title}.docx"
                # Clean filename
                filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
                
                document.file.save(
                    filename,
                    ContentFile(docx_content),
                    save=False
                )
                
                # Set current user for audit logging
                document.set_current_user(request.user)
                
                # Save the document
                document.save()
                
                # Convert to PDF if requested
                if form.cleaned_data.get('save_as_pdf'):
                    try:
                        pdf_conversion_success = document.convert_to_pdf(save_file=True)
                        if pdf_conversion_success:
                            document.save()  # Save the PDF field
                            messages.success(
                                request, 
                                _("Document created successfully and converted to PDF!")
                            )
                        else:
                            messages.warning(
                                request, 
                                _("Document created successfully, but PDF conversion failed. You can try converting it later from the document detail page.")
                            )
                    except Exception as e:
                        logger.error(f"PDF conversion error during document creation: {e}")
                        messages.warning(
                            request, 
                            _("Document created successfully, but PDF conversion failed. You can try converting it later from the document detail page.")
                        )
                else:
                    messages.success(request, _("Document created successfully!"))
                
                # Log document creation
                try:
                    AuditLog.log_action(
                        user=request.user,
                        document=document,
                        action='CREATE',
                        request=request,
                        creation_method='ckeditor',
                        file_size=document.file.size if document.file else 0,
                        editable=True,
                        pdf_generated=form.cleaned_data.get('save_as_pdf', False)
                    )
                except Exception as e:
                    logger.warning(f"Failed to log document creation: {e}")
                
                return redirect('documents:detail', pk=document.pk)
                
            except Exception as e:
                logger.error(f"Document creation failed: {str(e)}")
                messages.error(request, _("Document creation failed. Please try again."))
                return render(request, self.template_name, {'form': form})
        
        return render(request, self.template_name, {'form': form})

class DocumentUploadView(LoginRequiredMixin, CreateView):
    model = Document
    form_class = DocumentUploadForm
    template_name = 'documents/document_upload.html'
    success_url = reverse_lazy('documents:document_list')

    def dispatch(self, request, *args, **kwargs):
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to upload documents.")
            return redirect('documents:document_list')
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        import logging
        import os
        from django.utils import timezone
        
        logger = logging.getLogger(__name__)
        
        try:
            # Save document first to get an ID
            document = form.save(commit=False)
            document.uploaded_by = self.request.user
            document.set_current_user(self.request.user)  # For audit logging
            
            # Save with status as DRAFT initially
            document.status = Document.Status.DRAFT
            document.save()
            
            # Get file info for processing
            file_path = document.file.name
            file_size = document.file.size
            file_ext = os.path.splitext(file_path)[1].lower()
            
            logger.info(f"Processing upload: {file_path} ({file_size} bytes)")
            
            # Determine if this requires heavy processing
            requires_heavy_processing = (
                (file_ext in ['.mp3', '.wav', '.mp4', '.mov'] and file_size > 5 * 1024 * 1024) or  # Audio/Video > 5MB
                file_size > 50 * 1024 * 1024  # Any files larger than 50MB
            )
            
            # For small audio files, transcribe immediately (but only once)
            if (file_ext in ['.mp3', '.wav', '.mp4', '.mov'] and 
                file_size <= 5 * 1024 * 1024 and 
                not document.audio_transcript):
                logger.info("Small audio file detected, transcribing immediately")
                try:
                    document.document_type = Document.DocType.AUDIO if file_ext in ['.mp3', '.wav'] else Document.DocType.VIDEO
                    transcript = document.transcribe_media()
                    if transcript:
                        document.audio_transcript = transcript
                        logger.info(f"Immediate transcription completed: {len(transcript)} characters")
                    document.save()
                    requires_heavy_processing = False  # Skip background processing
                except Exception as e:
                    logger.error(f"Immediate transcription failed: {e}")
                    # Fall back to background processing
            
            if requires_heavy_processing:
                # For heavy processing, do basic extraction first
                logger.info("Heavy processing detected, scheduling background task")
                
                # Quick text extraction for immediate use
                if file_ext in ['.txt', '.md', '.html']:
                    try:
                        document.content = document._extract_plain_text()
                    except Exception as e:
                        logger.warning(f"Quick text extraction failed: {e}")
                elif file_ext in ['.docx', '.doc']:
                    try:
                        document.content = document._extract_docx_text()
                        # Try HTML extraction but don't fail if it doesn't work
                        try:
                            document.html_content = document.extract_html_content()
                        except Exception as e:
                            logger.warning(f"HTML extraction failed, will retry in background: {e}")
                    except Exception as e:
                        logger.warning(f"DOCX text extraction failed: {e}")
                elif file_ext == '.pdf':
                    try:
                        document.content = document._extract_pdf_text()
                    except Exception as e:
                        logger.warning(f"PDF text extraction failed: {e}")
                
                # Mark for background processing
                document.status = Document.Status.DRAFT
                document.save()
                
                # Log document creation
                try:
                    AuditLog.log_action(
                        user=self.request.user,
                        document=document,
                        action='CREATE',
                        request=self.request,
                        creation_method='upload',
                        file_size=file_size,
                        file_type=file_ext,
                        processing_type='background'
                    )
                except Exception as e:
                    logger.warning(f"Failed to log document creation: {e}")
                
                # Schedule background processing (you can implement this with Celery later)
                # For now, we'll do a simplified async approach
                self._schedule_background_processing(document)
                
                messages.info(
                    self.request, 
                    _("Document uploaded successfully! Content extraction and transcription are being processed in the background.")
                )
            else:
                # For smaller files, do full processing immediately
                logger.info("Light processing, extracting content immediately")
                
                # This will trigger the save method which handles extraction
                document.save()
                
                # Log document creation
                try:
                    AuditLog.log_action(
                        user=self.request.user,
                        document=document,
                        action='CREATE',
                        request=self.request,
                        creation_method='upload',
                        file_size=file_size,
                        file_type=file_ext,
                        processing_type='immediate'
                    )
                except Exception as e:
                    logger.warning(f"Failed to log document creation: {e}")
                
                messages.success(self.request, _("Document uploaded and processed successfully!"))
            
        except Exception as e:
            logger.error(f"Document upload failed: {str(e)}")
            # Create a more detailed error message
            error_msg = f"Upload failed: {str(e)}"
            
            # Try to clean up if document was partially saved
            if 'document' in locals() and hasattr(document, 'pk') and document.pk:
                try:
                    document.delete()
                except:
                    pass
            
            if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': error_msg})
            
            messages.error(self.request, _(f"Document upload failed: {error_msg}"))
            return self.form_invalid(form)
        
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'success': True,
                'document_id': document.pk,
                'processing_status': 'background' if requires_heavy_processing else 'complete'
            })
        
        return redirect(self.get_success_url())

    def _schedule_background_processing(self, document):
        """Schedule background processing for heavy operations"""
        import threading
        import logging
        
        logger = logging.getLogger(__name__)
        
        def background_process():
            try:
                logger.info(f"Starting background processing for document {document.pk}")
                
                # Reload document to avoid stale data
                document.refresh_from_db()
                
                # Process transcription for media files
                if (document.document_type in [Document.DocType.AUDIO, Document.DocType.VIDEO] and 
                    not document.audio_transcript):  # Only if not already transcribed
                    logger.info("Starting media transcription")
                    transcript = document.transcribe_media()
                    if transcript:
                        document.audio_transcript = transcript
                        logger.info(f"Transcription completed: {len(transcript)} characters")
                    else:
                        logger.warning("Transcription returned empty result")
                
                # Process HTML content for Word documents if not already done
                if (document.file.name.lower().endswith(('.docx', '.doc')) and 
                    not document.html_content):
                    logger.info("Extracting HTML content from Word document")
                    try:
                        html_content = document.extract_html_content()
                        if html_content:
                            document.html_content = html_content
                            logger.info("HTML content extraction completed")
                    except Exception as e:
                        logger.error(f"HTML extraction failed: {e}")
                
                # Update status and save
                document.status = Document.Status.REVIEW  # Ready for review
                document.save()
                
                logger.info(f"Background processing completed for document {document.pk}")
                
            except Exception as e:
                logger.error(f"Background processing failed for document {document.pk}: {e}")
                # Update status to indicate processing failed
                try:
                    document.refresh_from_db()
                    document.status = Document.Status.DRAFT  # Keep as draft with note
                    document.save()
                except:
                    pass
        
        # Start background thread
        thread = threading.Thread(target=background_process)
        thread.daemon = True
        thread.start()

    def form_invalid(self, form):
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({'success': False, 'errors': form.errors})
        return super().form_invalid(form)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs

class DocumentDetailView(LoginRequiredMixin, DetailView):
    model = Document
    template_name = 'documents/document_detail.html'

    def get_queryset(self):
        qs = super().get_queryset()
        if not self.request.user.has_perm('documents.can_view_confidential'):
            qs = qs.filter(confidential=False)
        return qs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        document = self.object
        search_query = self.request.GET.get('q')
        
        # Log document view
        try:
            AuditLog.log_action(
                user=self.request.user,
                document=document,
                action='VIEW',
                request=self.request,
                search_query=search_query if search_query else None
            )
        except Exception as e:
            logger.warning(f"Failed to log document view: {e}")
        
        # Determine content type and set flags
        context['is_media_file'] = document.document_type in [Document.DocType.AUDIO, Document.DocType.VIDEO]
        context['is_word_document'] = document.file.name.lower().endswith(('.docx', '.doc'))
        context['is_pdf_document'] = document.file.name.lower().endswith('.pdf')
        context['has_pdf_version'] = bool(document.pdf_file)
        context['can_convert_to_pdf'] = (
            context['is_word_document'] and 
            self.request.user.has_perm('documents.change_document')
        )
        
        # Set media type for audio/video files
        if context['is_media_file']:
            context['media_type'] = 'audio' if document.document_type == Document.DocType.AUDIO else 'video'
            context['display_content'] = document.audio_transcript
        elif context['is_word_document']:
            # For Word documents, prefer HTML content for better display
            if document.html_content:
                context['display_content'] = document.html_content
                context['content_type'] = 'html'
            else:
                # Try to extract HTML content
                try:
                    html_content = document.extract_html_content()
                    if html_content:
                        document.html_content = html_content
                        document.save(update_fields=['html_content'])
                        context['display_content'] = html_content
                        context['content_type'] = 'html'
                    else:
                        context['display_content'] = document.content
                        context['content_type'] = 'text'
                except Exception as e:
                    context['display_content'] = document.content
                    context['content_type'] = 'text'
                    context['extraction_error'] = str(e)
        else:
            context['display_content'] = document.content
            context['content_type'] = 'text'
        
        # Get document versions
        context['versions'] = document.versions.order_by('-version_number')
        
        # Add search highlighting if query exists
        if search_query:
            from .utils import get_search_queryset, extract_search_terms
            import json
            
            highlighted_doc = get_search_queryset(
                Document.objects.filter(pk=document.pk), 
                search_query
            ).first()
            
            search_terms = extract_search_terms(search_query)
            
            if highlighted_doc:
                context.update({
                    'title_headline': getattr(highlighted_doc, 'title_headline', None),
                    'content_headline': getattr(highlighted_doc, 'content_headline', None),
                    'transcript_headline': getattr(highlighted_doc, 'transcript_headline', None),
                    'search_query': search_query,
                    'search_terms': search_terms,
                    'search_terms_json': json.dumps(search_terms)
                })
            else:
                # Even if no highlights, still pass search terms for JavaScript highlighting
                context.update({
                    'search_query': search_query,
                    'search_terms': search_terms,
                    'search_terms_json': json.dumps(search_terms)
                })
        
        return context


class DocumentUpdateView(LoginRequiredMixin, UpdateView):
    model = Document
    form_class = DocumentUpdateForm
    template_name = 'documents/document_update.html'

    def dispatch(self, request, *args, **kwargs):
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to update documents.")
            return redirect('documents:document_list')
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse_lazy('documents:document_detail', kwargs={'pk': self.object.pk})

    def form_valid(self, form):
        old_status = self.object.status
        response = super().form_valid(form)
        
        # Log status changes and approvals
        if 'status' in form.changed_data:
            if form.instance.status == 'APPROVED':
                form.instance.reviewed_by = self.request.user
                AuditLog.log_action(
                    user=self.request.user,
                    document=self.object,
                    action='APPROVE',
                    request=self.request,
                    previous_status=old_status,
                    reviewed_by=self.request.user.get_full_name()
                )
            elif form.instance.status == 'REVIEW':
                AuditLog.log_action(
                    user=self.request.user,
                    document=self.object,
                    action='SUBMIT_REVIEW',
                    request=self.request,
                    previous_status=old_status
                )
            else:
                AuditLog.log_action(
                    user=self.request.user,
                    document=self.object,
                    action='STATUS_CHANGE',
                    request=self.request,
                    previous_status=old_status,
                    new_status=form.instance.status
                )
        
        # Log general updates for other fields
        if form.changed_data and 'status' not in form.changed_data:
            AuditLog.log_action(
                user=self.request.user,
                document=self.object,
                action='UPDATE',
                request=self.request,
                changed_fields=list(form.changed_data)
            )
        
        return response


class DocumentUpdateContentView(LoginRequiredMixin, View):
    def post(self, request, pk):
        document = get_object_or_404(Document, pk=pk)
        if not document.editable:
            messages.error(request, "This document type cannot be edited directly")
            return redirect('document_detail', pk=pk)

        old_content = document.content
        new_content = request.POST.get('content', '')
        document.content = new_content
        document.save()
        
        # Log content edit
        try:
            AuditLog.log_action(
                user=request.user,
                document=document,
                action='CONTENT_EDIT',
                request=request,
                content_length_before=len(old_content) if old_content else 0,
                content_length_after=len(new_content),
                content_changed=old_content != new_content
            )
        except Exception as e:
            logger.warning(f"Failed to log content edit: {e}")
        
        messages.success(request, "Document content updated successfully")
        return redirect('document_detail', pk=pk)

from django.shortcuts import get_object_or_404

class DocumentVersionCreateView(LoginRequiredMixin, CreateView):
    model = DocumentVersion
    form_class = DocumentVersionForm
    template_name = 'documents/version_create.html'

    def dispatch(self, request, *args, **kwargs):
        self.document = Document.objects.get(pk=self.kwargs['document_pk'])
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to create document versions.")
            return redirect('documents:document_detail', pk=self.document.pk)
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['document'] = self.document  # 🔥 Needed for template to work
        return context

    def form_valid(self, form):
        form.instance.document = self.document
        form.instance.created_by = self.request.user
        form.instance.version_number = self.document.versions.count() + 1
        response = super().form_valid(form)
        
        # Log version creation
        try:
            AuditLog.log_action(
                user=self.request.user,
                document=self.document,
                action='VERSION_CREATE',
                request=self.request,
                version_number=form.instance.version_number,
                changes=form.instance.changes
            )
        except Exception as e:
            logger.warning(f"Failed to log version creation: {e}")
        
        messages.success(self.request, _("New version created successfully!"))
        return response

    def get_success_url(self):
        return reverse_lazy('documents:document_detail', kwargs={'pk': self.document.pk})


from django.views.generic import ListView
from django.contrib import messages
from django.utils.translation import gettext_lazy as _
from django.utils import timezone

class DocumentArchiveView(LoginRequiredMixin, View):

    def post(self, request, pk):
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to archive documents.")
            return redirect('documents:document_detail', pk=pk)
            
        document = get_object_or_404(Document, pk=pk)
        
        try:
            with transaction.atomic():
                # Log archive action
                AuditLog.log_action(
                    user=request.user,
                    document=document,
                    action='ARCHIVE',
                    request=request,
                    original_status=document.get_status_display(),
                    confidential=document.confidential
                )
                
                # Update document
                document.is_archived = True
                document.archived_at = timezone.now()
                document.archived_by = request.user
                document.save()
                
        except Exception as e:
            logger.error(f"Document archive failed: {str(e)}")
            messages.error(request, _("Document archiving failed"))
            return redirect('documents:document_detail', pk=pk)

        messages.success(request, _("Document has been archived"))
        return redirect('documents:document_list')

class DocumentRestoreView(LoginRequiredMixin, View):

    def post(self, request, pk):
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to restore documents.")
            return redirect('documents:archived_documents')
            
        document = get_object_or_404(Document, pk=pk, is_archived=True)
        
        try:
            with transaction.atomic():
                # Log restore action
                AuditLog.log_action(
                    user=request.user,
                    document=document,
                    action='RESTORE',
                    request=request,
                    previous_archived_at=document.archived_at.isoformat() if document.archived_at else None,
                    archived_by=str(document.archived_by) if document.archived_by else None
                )
                
                document.is_archived = False
                document.archived_at = None
                document.archived_by = None
                document.save()
                
        except Exception as e:
            logger.error(f"Document restore failed: {str(e)}")
            messages.error(request, _("Document restoration failed"))
            return redirect('documents:archived_documents')

        messages.success(request, _("Document has been restored"))
        return redirect('documents:archived_documents')


class DocumentConvertToPDFView(LoginRequiredMixin, View):
    """Convert a Word document to PDF"""

    def post(self, request, pk):
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to convert documents.")
            return redirect('documents:document_detail', pk=pk)
            
        document = get_object_or_404(Document, pk=pk)
        
        # Check if user has permission to view this document
        if (document.confidential and 
            not request.user.has_perm('documents.can_view_confidential')):
            messages.error(request, _("Permission denied"))
            return redirect('documents:document_detail', pk=pk)
        
        # Check if document is a Word document
        if not document.file.name.lower().endswith(('.docx', '.doc')):
            messages.error(request, _("Only Word documents (.docx, .doc) can be converted to PDF"))
            return redirect('documents:document_detail', pk=pk)
        
        try:
            # Perform PDF conversion
            conversion_success = document.convert_to_pdf(save_file=True)
            
            if conversion_success:
                document.save()  # Save the PDF field
                
                # Log the conversion
                try:
                    AuditLog.log_action(
                        user=request.user,
                        document=document,
                        action='PDF_CONVERT',
                        request=request,
                        original_file=document.file.name,
                        pdf_file=document.pdf_file.name if document.pdf_file else None
                    )
                except Exception as e:
                    logger.warning(f"Failed to log PDF conversion: {e}")
                
                messages.success(request, _("Document successfully converted to PDF!"))
            else:
                messages.error(request, _("PDF conversion failed. Please check the document format and try again."))
                
        except Exception as e:
            logger.error(f"PDF conversion error: {e}")
            messages.error(request, _("PDF conversion failed due to an unexpected error."))
        
        return redirect('documents:document_detail', pk=pk)


class DocumentPDFDownloadView(LoginRequiredMixin, View):
    """Handle PDF downloads with audit logging"""
    
    def get(self, request, pk):
        document = get_object_or_404(Document, pk=pk)
        
        # Check if user has permission to view this document
        if (document.confidential and 
            not request.user.has_perm('documents.can_view_confidential')):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        # Check if PDF exists
        if not document.pdf_file:
            messages.error(request, _("PDF version not available for this document"))
            return redirect('documents:document_detail', pk=pk)
        
        # Log the download
        try:
            AuditLog.log_action(
                user=request.user,
                document=document,
                action='PDF_DOWNLOAD',
                request=request,
                file_name=document.pdf_file.name,
                file_size=document.pdf_file.size
            )
        except Exception as e:
            logger.warning(f"Failed to log PDF download: {e}")
        
        # Redirect to actual file URL
        from django.http import HttpResponseRedirect
        return HttpResponseRedirect(document.pdf_file.url)

class DocumentPermanentDeleteView(LoginRequiredMixin, DeleteView):
    model = Document
    template_name = 'documents/document_confirm_permanent_delete.html'
    success_url = reverse_lazy('documents:archived_documents')

    def dispatch(self, request, *args, **kwargs):
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to permanently delete documents.")
            return redirect('documents:archived_documents')
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        return super().get_queryset().filter(is_archived=True)

    def delete(self, request, *args, **kwargs):
        logger.info(f"DocumentPermanentDeleteView.delete called with args: {args}, kwargs: {kwargs}")
        document = self.get_object()
        logger.info(f"Retrieved document: ID={document.id}, Title={document.title}")
        
        # Store document info before deletion
        document_id = document.id
        document_title = document.title
        document_type = document.get_document_type_display()
        document_description = document.description
        document_status = document.get_status_display()
        document_confidential = document.confidential
        document_archived_at = document.archived_at.isoformat() if document.archived_at else None
        
        logger.info(f"Stored document info: ID={document_id}, Title={document_title}, Type={document_type}")
        
        try:
            # Create audit log entry manually with stored data (before deletion)
            from django.utils import timezone
            
            logger.info(f"Creating audit log for permanent deletion of document {document_id}: {document_title}")
            
            # Get IP address and user agent from request
            ip_address = None
            user_agent = ''
            session_key = None
            
            if request:
                x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
                if x_forwarded_for:
                    ip_address = x_forwarded_for.split(',')[0]
                else:
                    ip_address = request.META.get('REMOTE_ADDR')
                user_agent = request.META.get('HTTP_USER_AGENT', '')[:500]
                if hasattr(request, 'session') and request.session.session_key:
                    session_key = request.session.session_key
            
            logger.info(f"Request info: IP={ip_address}, User={request.user.username if request.user else 'None'}")
            
            # Create the audit log entry directly
            audit_log = AuditLog.objects.create(
                user=request.user,
                document_id=document_id,
                document_title=document_title,
                document_type=document_type,
                action='DELETE',
                additional_info={
                    'original_description': document_description,
                    'original_status': document_status,
                    'confidential': document_confidential,
                    'archived_at': document_archived_at,
                    'delete_method': 'permanent_delete_view'
                },
                ip_address=ip_address,
                user_agent=user_agent,
                session_key=session_key
            )
            
            logger.info(f"✅ Audit log created successfully with ID: {audit_log.id}")
            
            # Verify the audit log was actually saved
            try:
                saved_log = AuditLog.objects.get(id=audit_log.id)
                logger.info(f"✅ Audit log verification: Found log with ID {saved_log.id}, action={saved_log.action}, document_id={saved_log.document_id}")
            except AuditLog.DoesNotExist:
                logger.error("❌ Audit log was not saved properly!")
            
            # Now delete the document
            logger.info(f"Proceeding to delete document {document_id}")
            response = super().delete(request, *args, **kwargs)
            logger.info(f"✅ Document {document_id} deleted successfully")
            
            # Verify the audit log still exists after document deletion
            try:
                still_exists = AuditLog.objects.get(id=audit_log.id)
                logger.info(f"✅ Audit log still exists after document deletion: ID {still_exists.id}")
            except AuditLog.DoesNotExist:
                logger.error(f"❌ Audit log with ID {audit_log.id} was deleted along with the document!")
            
            messages.success(request, _("Document has been permanently deleted"))
            return response
                
        except Exception as e:
            logger.error(f"❌ Permanent delete failed: {str(e)}")
            logger.exception("Full exception details:")
            messages.error(request, _("Permanent deletion failed"))
            return redirect('documents:archived_documents')

class ArchivedDocumentListView(LoginRequiredMixin, ListView):
    model = Document
    template_name = 'documents/archived_document_list.html'
    context_object_name = 'documents'

    def dispatch(self, request, *args, **kwargs):
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to view archived documents.")
            return redirect('documents:document_list')
        return super().dispatch(request, *args, **kwargs)
    paginate_by = 20

    def get_queryset(self):
        return super().get_queryset().filter(is_archived=True).order_by('-archived_at')

class AuditLogListView(LoginRequiredMixin, ListView):
    model = AuditLog
    template_name = 'documents/audit_log.html'
    paginate_by = 20
    ordering = ['-timestamp']
    
    def get_queryset(self):
        queryset = super().get_queryset()
        
        # Filter by action if provided
        action_filter = self.request.GET.get('action')
        if action_filter:
            queryset = queryset.filter(action=action_filter)
        
        # Filter by date range if provided
        date_from = self.request.GET.get('date_from')
        date_to = self.request.GET.get('date_to')
        
        if date_from:
            queryset = queryset.filter(timestamp__gte=date_from)
        if date_to:
            queryset = queryset.filter(timestamp__lte=date_to)
        
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Calculate statistics
        from django.utils import timezone
        from django.db.models import Count
        
        today = timezone.now().date()
        all_logs = AuditLog.objects.all()
        
        context.update({
            'total_logs': all_logs.count(),
            'today_logs': all_logs.filter(timestamp__date=today).count(),
            'unique_users': all_logs.values('user').distinct().count(),
            'unique_documents': all_logs.values('document_id').distinct().count(),
        })
        
        return context

class DocumentWordEditView(LoginRequiredMixin, UpdateView):
    """View for editing Word documents with CKEditor"""
    model = Document
    form_class = DocumentWordEditForm
    template_name = 'documents/document_word_edit.html'

    def dispatch(self, request, *args, **kwargs):
        if not can_manage_documents(request.user):
            messages.error(request, "You don't have permission to edit documents.")
            return redirect('documents:document_list')
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        qs = super().get_queryset()
        # Only allow editing of Word documents
        return qs.filter(file__iendswith='.docx') | qs.filter(file__iendswith='.doc')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        document = self.object
        
        # Ensure HTML content is available
        if not document.html_content and document.file:
            try:
                html_content = document.extract_html_content()
                if html_content:
                    document.html_content = html_content
                    document.save(update_fields=['html_content'])
                else:
                    # Create basic HTML from text content
                    if document.content:
                        paragraphs = document.content.split('\n\n')
                        html_content = ''.join([f"<p>{para}</p>" for para in paragraphs if para.strip()])
                        document.html_content = html_content
                        document.save(update_fields=['html_content'])
            except Exception as e:
                context['extraction_error'] = str(e)
                logger.error(f"Failed to extract HTML content: {e}")
        
        return context

    def form_valid(self, form):
        # Create version before updating if this is a significant change
        old_html_content = self.object.html_content
        content_changed = self.object.html_content != form.cleaned_data.get('html_content')
        
        if content_changed:
            try:
                # Create a new version
                version = DocumentVersion.objects.create(
                    document=self.object,
                    file=self.object.file,
                    version_number=self.object.versions.count() + 1,
                    changes="Content edited via rich text editor",
                    created_by=self.request.user
                )
            except Exception as e:
                logger.warning(f"Could not create version: {e}")
        
        try:
            response = super().form_valid(form)
            
            # Log the word document edit
            if content_changed:
                try:
                    AuditLog.log_action(
                        user=self.request.user,
                        document=self.object,
                        action='CONTENT_EDIT',
                        request=self.request,
                        edit_type='rich_text_editor',
                        html_content_length_before=len(old_html_content) if old_html_content else 0,
                        html_content_length_after=len(form.cleaned_data.get('html_content', '')),
                        version_created=True
                    )
                except Exception as e:
                    logger.warning(f"Failed to log word document edit: {e}")
            
            messages.success(self.request, _("Document updated successfully!"))
            return response
        except Exception as e:
            logger.error(f"Error saving document: {e}")
            # Add error message to form
            form.add_error(None, _("There was an error saving the document. Please try again."))
            return self.form_invalid(form)

    def get_success_url(self):
        return reverse_lazy('documents:document_detail', kwargs={'pk': self.object.pk})
    

class PendingReviewListView(LoginRequiredMixin, ListView):
    """View for documents pending review"""
    model = Document
    template_name = 'documents/pending_review.html'
    context_object_name = 'documents'
    paginate_by = 20

    def get_queryset(self):
        qs = super().get_queryset().filter(status='REVIEW', is_archived=False)
        
        if not self.request.user.has_perm('documents.can_view_confidential'):
            qs = qs.filter(confidential=False)
            
        return qs.order_by('-uploaded_at')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = 'Documents Pending Review'
        context['total_pending'] = self.get_queryset().count()
        return context


class AuditLogExportView(LoginRequiredMixin, View):
    """Export audit logs in various formats"""
    
    def get(self, request):
        format_type = request.GET.get('format', 'csv')
        
        # Get filtered queryset
        queryset = AuditLog.objects.all().order_by('-timestamp')
        
        # Apply filters
        action_filter = request.GET.get('action')
        if action_filter:
            queryset = queryset.filter(action=action_filter)
        
        date_from = request.GET.get('date_from')
        date_to = request.GET.get('date_to')
        
        if date_from:
            queryset = queryset.filter(timestamp__gte=date_from)
        if date_to:
            queryset = queryset.filter(timestamp__lte=date_to)
        
        # Limit to recent logs for performance
        queryset = queryset[:1000]
        
        if format_type == 'csv':
            return self._export_csv(queryset)
        elif format_type == 'pdf':
            return self._export_pdf(queryset)
        elif format_type == 'docx':
            return self._export_docx(queryset)
        else:
            return JsonResponse({'error': 'Invalid format'}, status=400)
    
    def _export_csv(self, queryset):
        import csv
        from django.http import HttpResponse
        from django.utils import timezone
        
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="audit_log_{timezone.now().strftime("%Y%m%d_%H%M%S")}.csv"'
        
        writer = csv.writer(response)
        writer.writerow([
            'Timestamp', 'User', 'User Role', 'Action', 'Document ID', 'Document Title', 
            'Document Type', 'IP Address', 'Additional Details'
        ])
        
        for log in queryset:
            additional_details = '; '.join([f"{k}: {v}" for k, v in log.additional_info.items()])
            writer.writerow([
                log.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
                log.user.get_full_name() if log.user else 'Anonymous',
                log.user.get_role_display() if log.user else '',
                log.get_action_display(),
                log.document_id,
                log.document_title,
                log.document_type,
                log.ip_address or '',
                additional_details
            ])
        
        return response
    
    def _export_pdf(self, queryset):
        from django.http import HttpResponse
        from django.template.loader import render_to_string
        from django.utils import timezone
        import tempfile
        
        # For PDF generation, you'd typically use libraries like ReportLab or weasyprint
        # For now, I'll create a simple HTML response that can be printed as PDF
        
        context = {
            'logs': queryset,
            'generated_at': timezone.now(),
            'total_count': queryset.count(),
        }
        
        html_content = render_to_string('documents/audit_log_pdf.html', context)
        
        response = HttpResponse(html_content, content_type='text/html')
        response['Content-Disposition'] = f'inline; filename="audit_log_{timezone.now().strftime("%Y%m%d_%H%M%S")}.html"'
        
        return response
    
    def _export_docx(self, queryset):
        from django.http import HttpResponse
        from django.utils import timezone
        import tempfile
        import os
        
        try:
            from docx import Document
            from docx.shared import Inches
            
            # Create document
            doc = Document()
            doc.add_heading('Audit Log Report', 0)
            
            # Add metadata
            doc.add_paragraph(f'Generated on: {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Total entries: {queryset.count()}')
            doc.add_paragraph('')
            
            # Create table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Timestamp'
            hdr_cells[1].text = 'User'
            hdr_cells[2].text = 'Action'
            hdr_cells[3].text = 'Document'
            hdr_cells[4].text = 'Details'
            hdr_cells[5].text = 'IP Address'
            
            # Add data
            for log in queryset[:500]:  # Limit for performance
                row_cells = table.add_row().cells
                row_cells[0].text = log.timestamp.strftime('%Y-%m-%d %H:%M')
                row_cells[1].text = log.user.get_full_name() if log.user else 'Anonymous'
                row_cells[2].text = log.get_action_display()
                row_cells[3].text = f"{log.document_title} (ID: {log.document_id})"
                row_cells[4].text = '; '.join([f"{k}: {v}" for k, v in log.additional_info.items()])[:100]
                row_cells[5].text = log.ip_address or ''
            
            # Save to memory
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            # Return file
            with open(temp_file.name, 'rb') as f:
                response = HttpResponse(
                    f.read(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="audit_log_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx"'
            
            # Clean up
            os.unlink(temp_file.name)
            
            return response
            
        except ImportError:
            # Fallback to simple text file
            response = HttpResponse(content_type='text/plain')
            response['Content-Disposition'] = f'attachment; filename="audit_log_{timezone.now().strftime("%Y%m%d_%H%M%S")}.txt"'
            
            response.write('AUDIT LOG REPORT\n')
            response.write('=' * 50 + '\n\n')
            response.write(f'Generated on: {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
            response.write(f'Total entries: {queryset.count()}\n\n')
            
            for log in queryset[:500]:
                response.write(f'{log.timestamp.strftime("%Y-%m-%d %H:%M:%S")} | ')
                response.write(f'{log.user.get_full_name() if log.user else "Anonymous"} | ')
                response.write(f'{log.get_action_display()} | ')
                response.write(f'{log.document_title} (ID: {log.document_id}) | ')
                response.write(f'{log.ip_address or ""}\n')
                
                if log.additional_info:
                    for k, v in log.additional_info.items():
                        response.write(f'  {k}: {v}\n')
                response.write('\n')
            
            return response


class AuditLogDetailView(LoginRequiredMixin, View):
    """Get detailed information about a specific audit log entry"""
    
    def get(self, request, log_id):
        try:
            log = AuditLog.objects.get(id=log_id)
            
            # Build detailed HTML
            html_content = f"""
            <div class="row">
                <div class="col-md-6">
                    <h6>Basic Information</h6>
                    <table class="table table-sm">
                        <tr><th>Timestamp:</th><td>{log.timestamp.strftime('%Y-%m-%d %H:%M:%S %Z')}</td></tr>
                        <tr><th>User:</th><td>{log.user.get_full_name() if log.user else 'Anonymous'}</td></tr>
                        <tr><th>User Role:</th><td>{log.user.get_role_display() if log.user else 'N/A'}</td></tr>
                        <tr><th>Action:</th><td>{log.get_action_display()}</td></tr>
                        <tr><th>IP Address:</th><td>{log.ip_address or 'N/A'}</td></tr>
                        <tr><th>Session:</th><td>{log.session_key or 'N/A'}</td></tr>
                    </table>
                </div>
                <div class="col-md-6">
                    <h6>Document Information</h6>
                    <table class="table table-sm">
                        <tr><th>Document ID:</th><td>{log.document_id}</td></tr>
                        <tr><th>Title:</th><td>{log.document_title}</td></tr>
                        <tr><th>Type:</th><td>{log.document_type}</td></tr>
                    </table>
                </div>
            </div>
            """
            
            if log.additional_info:
                html_content += """
                <div class="row mt-3">
                    <div class="col-12">
                        <h6>Additional Details</h6>
                        <div class="table-responsive">
                            <table class="table table-sm table-striped">
                """
                for key, value in log.additional_info.items():
                    html_content += f"<tr><th>{key.replace('_', ' ').title()}:</th><td>{value}</td></tr>"
                
                html_content += "</table></div></div></div>"
            
            if log.user_agent:
                html_content += f"""
                <div class="row mt-3">
                    <div class="col-12">
                        <h6>Browser Information</h6>
                        <small class="text-muted">{log.user_agent}</small>
                    </div>
                </div>
                """
            
            return JsonResponse({'html': html_content})
            
        except AuditLog.DoesNotExist:
            return JsonResponse({'error': 'Log entry not found'}, status=404)