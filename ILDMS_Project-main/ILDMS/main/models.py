from django.db import models
# documents/models.py
import os
import uuid
import tempfile
import logging
from django.contrib.postgres.indexes import GinIndex
from django.contrib.auth.models import AbstractUser, Group, Permission
from django.db import models
from django.contrib.auth.models import AbstractUser, Group,BaseUserManager
from django.utils import timezone
from django.utils.timezone import now
from django.utils.translation import gettext_lazy as _
from django.core.validators import FileExtensionValidator, RegexValidator
from django.core.exceptions import ValidationError
from django.contrib.postgres.search import SearchVector, SearchVectorField
from django.contrib.postgres.indexes import GinIndex
import re

# Security imports for file validation
try:
    import magic
except ImportError:
    magic = None

# Optional imports for content processing
try:
    import whisper
except ImportError:
    whisper = None


def validate_safe_filename(filename):
    """Validate filename for security"""
    if not filename:
        raise ValidationError(_('Filename cannot be empty.'))
    
    # Check for dangerous patterns
    dangerous_patterns = ['../', '..\\', './', '.\\', '~/', '~\\']
    filename_lower = filename.lower()
    
    for pattern in dangerous_patterns:
        if pattern in filename_lower:
            raise ValidationError(_('Filename contains invalid path characters.'))
    
    # Check for null bytes and control characters
    if '\x00' in filename or any(ord(c) < 32 for c in filename if c not in '\t\n\r'):
        raise ValidationError(_('Filename contains invalid characters.'))
    
    # Check length
    if len(filename) > 255:
        raise ValidationError(_('Filename too long (max 255 characters).'))


def validate_title_length(title):
    """Validate title length and content"""
    if not title or len(title.strip()) < 3:
        raise ValidationError(_('Title must be at least 3 characters long.'))
    
    if len(title) > 255:
        raise ValidationError(_('Title must be less than 255 characters.'))


def validate_description_length(description):
    """Validate description length"""
    if description and len(description) > 2000:
        raise ValidationError(_('Description must be less than 2000 characters.'))


def validate_secure_content(content):
    """Validate content for XSS and other threats"""
    if not content:
        return
    
    # Check for script injection attempts
    dangerous_patterns = [
        r'<script[^>]*>.*?</script>',
        r'javascript:',
        r'vbscript:',
        r'onload\s*=',
        r'onerror\s*='
    ]
    
    content_lower = content.lower()
    for pattern in dangerous_patterns:
        if re.search(pattern, content_lower, re.IGNORECASE | re.DOTALL):
            raise ValidationError(_('Content contains potentially dangerous script elements.'))


# Enhanced file extension validator
ALLOWED_EXTENSIONS = [
    'pdf', 'docx', 'doc', 'ppt', 'pptx', 'xls', 'xlsx', 
    'mp3', 'wav', 'mp4', 'mov', 'txt', 'md'
]

secure_file_validator = FileExtensionValidator(
    allowed_extensions=ALLOWED_EXTENSIONS,
    message=_('File type not allowed. Allowed types: %(allowed_extensions)s')
)

class CustomUserManager(BaseUserManager):
    def create_user(self, username, email, password=None, **extra_fields):
        """
        Creates and saves a User with the given email and password.
        """
        if not email:
            raise ValueError('Users must have an email address')

        if not username:
            raise ValueError('Users must have a username')

        email = self.normalize_email(email)
        user = self.model(
            username=username,
            email=email,
            **extra_fields
        )
        user.set_password(password)
        user.save(using=self._db)
        return user

    def create_superuser(self, username, email, password, **extra_fields):
        """
        Creates and saves a superuser with the given email and password.
        """
        extra_fields.setdefault('is_staff', True)
        extra_fields.setdefault('is_superuser', True)
        extra_fields.setdefault('is_active', True)
        extra_fields.setdefault('role', 'LAWYER')

        if extra_fields.get('is_staff') is not True:
            raise ValueError('Superuser must have is_staff=True.')
        if extra_fields.get('is_superuser') is not True:
            raise ValueError('Superuser must have is_superuser=True.')

        return self.create_user(username, email, password, **extra_fields)


class CustomUserManager(BaseUserManager):
    def create_user(self, username, email, password=None, **extra_fields):
        """
        Creates and saves a User with the given email and password.
        """
        if not email:
            raise ValueError('Users must have an email address')

        if not username:
            raise ValueError('Users must have a username')

        email = self.normalize_email(email)
        user = self.model(
            username=username,
            email=email,
            **extra_fields
        )
        user.set_password(password)
        user.save(using=self._db)
        return user

    def create_superuser(self, username, email, password, **extra_fields):
        """
        Creates and saves a superuser with the given email and password.
        """
        extra_fields.setdefault('is_staff', True)
        extra_fields.setdefault('is_superuser', True)
        extra_fields.setdefault('is_active', True)
        extra_fields.setdefault('role', User.Role.LAWYER)
        extra_fields.setdefault('is_verified', True)

        if extra_fields.get('is_staff') is not True:
            raise ValueError('Superuser must have is_staff=True.')
        if extra_fields.get('is_superuser') is not True:
            raise ValueError('Superuser must have is_superuser=True.')

        return self.create_user(username, email, password, **extra_fields)


class User(AbstractUser):
    class Role(models.TextChoices):
        LAWYER = 'LAWYER', _('Lawyer')
        SENIOR_PARALEGAL = 'SENIOR_PARALEGAL', _('Senior Paralegal')
        PARALEGAL = 'PARALEGAL', _('Paralegal')
        CLIENT = 'CLIENT', _('Client')
        ADMIN = 'ADMIN', _('System Admin')

    role = models.CharField(
        max_length=20,
        choices=Role.choices,
        default=Role.CLIENT,
        verbose_name=_('User Role')
    )

    phone = models.CharField(
        max_length=20,
        blank=True,
        validators=[
            RegexValidator(
                regex=r'^\+?1?\d{9,15}$',
                message=_("Phone number must be entered in the format: '+999999999'. Up to 15 digits allowed.")
            )
        ],
        verbose_name=_('Phone Number')
    )

    is_verified = models.BooleanField(
        default=False,
        verbose_name=_('Email Verified')
    )

    last_updated = models.DateTimeField(
        auto_now=True,
        verbose_name=_('Last Updated')
    )

    # Additional fields for legal practice
    bar_number = models.CharField(
        max_length=50,
        blank=True,
        null=True,
        verbose_name=_('Bar Number'),
        help_text=_('Required for lawyers')
    )

    specialization = models.CharField(
        max_length=100,
        blank=True,
        verbose_name=_('Legal Specialization')
    )

    last_login_ip = models.GenericIPAddressField(
        blank=True,
        null=True,
        verbose_name=_('Last Login IP')
    )


    # Metadata
    class Meta:
        verbose_name = _('User')
        verbose_name_plural = _('Users')
        ordering = ['-date_joined']
        permissions = [
            ('can_view_confidential', _('Can view confidential documents')),
            ('can_approve_documents', _('Can approve legal documents')),
            ('can_manage_users', _('Can manage other users')),
        ]

    def __str__(self):
        name = self.get_full_name() or self.username
        return f"{name} ({self.get_role_display()})"

    def clean(self):
        super().clean()
        if self.role == self.Role.LAWYER and not self.bar_number:
            raise ValidationError({
                'bar_number': _('Lawyers must provide a valid bar number')
            })

    def save(self, *args, **kwargs):
        # Automatically set is_staff for certain roles
        if self.role in [self.Role.LAWYER, self.Role.ADMIN, self.Role.SENIOR_PARALEGAL]:
            self.is_staff = True
        super().save(*args, **kwargs)

    def get_role_permissions(self):
        """Return permissions automatically granted by the user's role"""
        if self.role == self.Role.ADMIN:
            return Permission.objects.all()
        elif self.role == self.Role.LAWYER:
            return Permission.objects.filter(
                codename__in=['can_view_confidential', 'can_approve_documents']
            )
        elif self.role == self.Role.SENIOR_PARALEGAL:
            return Permission.objects.filter(
                codename='can_view_confidential'
            )
        return Permission.objects.none()

    @property
    def is_lawyer(self):
        return self.role == self.Role.LAWYER

    @property
    def is_paralegal(self):
        return self.role in [self.Role.PARALEGAL, self.Role.SENIOR_PARALEGAL]

    @property
    def can_approve_docs(self):
        return self.has_perm('main.can_approve_documents')

    objects = CustomUserManager()

# documents

def document_upload_path(instance, filename):
    return f"documents/user_{instance.uploaded_by.id}/{timezone.now().strftime('%Y%m%d')}/{filename}"


class Document(models.Model):
    class DocType(models.TextChoices):
        CONSTITUTION = 'CONSTITUTION', _('Constitution')
        BILL = 'BILL', _('Bill')
        ACT = 'ACT', _('Act of Parliament')
        LEGAL_NOTICE = 'LEGAL_NOTICE', _('Legal Notice')
        PRACTICE_NOTE = 'PRACTICE_NOTE', _('Practice Note')
        TREATY = 'TREATY', _('Treaty')
        CASE_LAW = 'CASE_LAW', _('Case Law / Judgment')
        CONTRACT = 'CONTRACT', _('Contract / Agreement')
        COURT_FILING = 'COURT_FILING', _('Court Filing')
        CORRESPONDENCE = 'CORRESPONDENCE', _('Correspondence')
        RESEARCH = 'RESEARCH', _('Legal Research')
        AUDIO = 'AUDIO', _('Audio Recording')
        VIDEO = 'VIDEO', _('Video Recording')
        OTHER = 'OTHER', _('Other')

    class Status(models.TextChoices):
        DRAFT = 'DRAFT', _('Draft')
        REVIEW = 'REVIEW', _('Under Review')
        APPROVED = 'APPROVED', _('Approved')
        ARCHIVED = 'ARCHIVED', _('Archived')

    title = models.CharField(
        max_length=255, 
        verbose_name=_('Title'),
        validators=[validate_title_length],
        help_text=_('Document title (3-255 characters)')
    )
    document_type = models.CharField(
        max_length=20,
        choices=DocType.choices,
        default=DocType.OTHER,
        verbose_name=_('Document Type')
    )
    file = models.FileField(
        upload_to=document_upload_path,
        validators=[secure_file_validator],
        verbose_name=_('File'),
        help_text=_('Allowed file types: PDF, DOCX, DOC, PPT, XLS, MP3, WAV, MP4, MOV, TXT, MD')
    )
    pdf_file = models.FileField(
        upload_to='documents/pdf/',
        blank=True,
        null=True,
        validators=[FileExtensionValidator(allowed_extensions=['pdf'])],
        verbose_name=_('PDF Version'),
        help_text=_('Auto-generated PDF version of the document')
    )
    content = models.TextField(
        blank=True, 
        verbose_name=_('Extracted Content'),
        validators=[validate_secure_content]
    )
    html_content = models.TextField(
        blank=True, 
        verbose_name=_('HTML Content'), 
        help_text=_('Rich text content for display and editing'),
        validators=[validate_secure_content]
    )
    audio_transcript = models.TextField(
        blank=True, 
        verbose_name=_('Audio Transcript'),
        help_text=_('Transcript of audio/video content')
    )
    description = models.TextField(
        blank=True, 
        verbose_name=_('Description'),
        validators=[validate_description_length],
        help_text=_('Brief description of the document (max 2000 characters)')
    )
    confidential = models.BooleanField(default=False, verbose_name=_('Confidential'))
    status = models.CharField(
        max_length=20,
        choices=Status.choices,
        default=Status.DRAFT,
        verbose_name=_('Status')
    )
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.PROTECT,
        related_name='uploaded_documents',
        verbose_name=_('Uploaded By')
    )
    uploaded_at = models.DateTimeField(auto_now_add=True, verbose_name=_('Uploaded At'))
    modified_at = models.DateTimeField(auto_now=True, verbose_name=_('Last Modified'))
    tags = models.ManyToManyField('Tag', blank=True, verbose_name=_('Tags'))
    editable = models.BooleanField(
        default=False,
        verbose_name=_('Editable'),
        help_text=_('Can this document be edited directly in the system?')
    )
    reviewed_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='reviewed_documents',
        verbose_name=_('Reviewed By')
    )
    search_vector = SearchVectorField(null=True, editable=False)

    is_archived = models.BooleanField(
        default=False,
        verbose_name=_('Archived'),
        help_text=_('Whether the document has been archived/deleted')
    )
    archived_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name=_('Archived At')
    )
    archived_by = models.ForeignKey(
        User,
        null=True,
        blank=True,
        on_delete=models.SET_NULL,
        related_name='archived_documents',
        verbose_name=_('Archived By')
    )

    class Meta:
        ordering = ['-uploaded_at']
        permissions = [
            ("can_analyze_document", _("Can perform NLP analysis")),
            ("can_share_document", _("Can share documents with others")),
            ("can_mark_confidential", _("Can mark documents as confidential")),
            ("can_approve_document", _("Can approve documents")),
            ("can_delete_document", _("Can delete documents")),
            ("can_view_archived", _("Can view archived documents")),
            ("can_permanently_delete", _("Can permanently delete documents")),
        ]
        verbose_name = _('Document')
        verbose_name_plural = _('Documents')
        indexes = [
            GinIndex(fields=['search_vector']),
            models.Index(fields=['title']),
            models.Index(fields=['document_type']),
            models.Index(fields=['status']),
            models.Index(fields=['uploaded_at']),
        ]

    def __str__(self):
        return f"{self.title} ({self.get_document_type_display()})"

    def clean(self):
        super().clean()
        
        # Additional security validations
        if self.confidential and not hasattr(self, 'uploaded_by'):
            raise ValidationError(_("Confidential documents must have an uploader."))
        
        # Validate file extension against MIME type if file exists
        if self.file:
            validate_safe_filename(self.file.name)
            
            # Additional file size checks
            if self.file.size > 500 * 1024 * 1024:  # 500MB absolute limit
                raise ValidationError(_("File size exceeds maximum limit of 500MB."))
        
        # Validate title content
        if self.title:
            # Remove potentially dangerous characters
            cleaned_title = re.sub(r'[<>:"/\\|?*]', '_', self.title)
            if cleaned_title != self.title:
                self.title = cleaned_title
        
        # Validate description content
        if self.description:
            if len(self.description) > 2000:
                raise ValidationError(_("Description exceeds maximum length of 2000 characters."))

    def get_file_size_display(self):
        """Return human-readable file size"""
        if not self.file:
            return "No file"
        
        size = self.file.size
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"

    def is_media_file(self):
        """Check if file is audio or video"""
        if not self.file:
            return False
        ext = self.file.name.lower().split('.')[-1]
        return ext in ['mp3', 'wav', 'mp4', 'mov']

    def is_office_document(self):
        """Check if file is an Office document"""
        if not self.file:
            return False
        ext = self.file.name.lower().split('.')[-1]
        return ext in ['docx', 'doc', 'pptx', 'ppt', 'xlsx', 'xls']

    def get_file_icon(self):
        """Get appropriate icon for file type"""
        if not self.file:
            return 'file'
        
        ext = self.file.name.lower().split('.')[-1]
        icon_map = {
            'pdf': 'file-pdf',
            'docx': 'file-word', 'doc': 'file-word',
            'pptx': 'file-powerpoint', 'ppt': 'file-powerpoint',
            'xlsx': 'file-excel', 'xls': 'file-excel',
            'mp3': 'file-audio', 'wav': 'file-audio',
            'mp4': 'file-video', 'mov': 'file-video',
            'txt': 'file-text', 'md': 'file-text',
        }
        return icon_map.get(ext, 'file')

    def clean(self):
        super().clean()
        if self.confidential and not hasattr(self, 'uploaded_by'):
            raise ValidationError(_("Confidential documents must have an uploader."))

    def save(self, *args, **kwargs):
        import logging
        logger = logging.getLogger(__name__)
        
        is_new_file = not self.pk or 'file' in kwargs.get('update_fields', [])
        is_new_document = not self.pk
        
        # Only process content on new file uploads
        if is_new_file and self.file:
            logger.info(f"Processing new file: {self.file.name}")
            
            try:
                # Extract text content
                self.content = self.extract_text()
                logger.info(f"Text extraction completed: {len(self.content or '')} characters")
                
                # Generate HTML content for Word documents (lighter operation)
                if self.file.name.lower().endswith(('.docx', '.doc')):
                    try:
                        self.html_content = self.extract_html_content()
                        logger.info(f"HTML extraction completed: {len(self.html_content or '')} characters")
                    except Exception as e:
                        logger.warning(f"HTML extraction failed: {e}")
                        self.html_content = ""
                
                # Set document type based on file extension if not already set
                if self.document_type == self.DocType.OTHER:
                    self._auto_detect_document_type()
                
                # Transcription will be handled by background processing for efficiency
                # Do NOT do it in save() method to avoid duplication
                # The upload view will handle transcription via background processing
                
                # Set editable flag
                self.editable = self.file.name.lower().endswith(('.txt', '.md', '.html', '.docx', '.doc'))
                
            except Exception as e:
                logger.error(f"Content processing failed: {e}")
                # Don't fail the save, just log the error
                self.content = f"Content extraction failed: {str(e)}"

        super().save(*args, **kwargs)

        # Update search vector after save
        if 'update_fields' not in kwargs or 'search_vector' not in kwargs['update_fields']:
            try:
                self.update_search_vector()
            except Exception as e:
                logger.warning(f"Search vector update failed: {e}")
        
        # Create audit log entry
        try:
            # Import here to avoid circular imports
            from django.contrib.auth import get_user_model
            
            # Try to get current user from thread local storage or request context
            # This is a simplified approach - in production you'd want a more robust solution
            current_user = getattr(self, '_current_user', None)
            
            if current_user:
                action = 'CREATE' if is_new_document else 'UPDATE'
                additional_info = {}
                
                if is_new_document:
                    additional_info.update({
                        'file_size': self.file.size if self.file else 0,
                        'file_type': self.document_type,
                        'has_content': bool(self.content),
                    })
                else:
                    # Track what was updated
                    changed_fields = getattr(self, '_changed_fields', [])
                    if changed_fields:
                        additional_info['changed_fields'] = changed_fields
                
                AuditLog.log_action(
                    user=current_user,
                    document=self,
                    action=action,
                    **additional_info
                )
        except Exception as e:
            logger.warning(f"Failed to create audit log: {e}")

    def set_current_user(self, user):
        """Helper method to set current user for audit logging"""
        self._current_user = user

    def _auto_detect_document_type(self):
        """Auto-detect document type based on file extension and content"""
        file_ext = self.file.name.lower().split('.')[-1] if self.file.name else ''
        
        type_mapping = {
            'mp3': self.DocType.AUDIO,
            'wav': self.DocType.AUDIO,
            'mp4': self.DocType.VIDEO,
            'mov': self.DocType.VIDEO,
            'avi': self.DocType.VIDEO,
        }
        
        if file_ext in type_mapping:
            self.document_type = type_mapping[file_ext]

    def update_search_vector(self):
        """Update search vector with comprehensive field coverage"""
        try:
            Document.objects.filter(pk=self.pk).update(
                search_vector=(
                    SearchVector('title', weight='A', config='english') +
                    SearchVector('content', weight='B', config='english') +
                    SearchVector('html_content', weight='B', config='english') +
                    SearchVector('description', weight='C', config='english') +
                    SearchVector('audio_transcript', weight='D', config='english')
                )
            )
            logger = logging.getLogger(__name__)
            logger.info(f"Updated search vector for document {self.pk}: {self.title}")
        except Exception as e:
            logger = logging.getLogger(__name__)
            logger.error(f"Failed to update search vector for document {self.pk}: {e}")

    def transcribe_media(self):
        """Transcribe audio/video files using Whisper with improved efficiency"""
        import logging
        import tempfile
        import os
        
        logger = logging.getLogger(__name__)
        
        if not self.file:
            return ""
            
        try:
            file_ext = self.file.name.lower().split('.')[-1] if self.file.name else ''
            
            if file_ext not in ['mp3', 'wav', 'mp4', 'mov', 'avi', 'm4a', 'aac']:
                logger.info(f"File extension .{file_ext} not supported for transcription")
                return ""
            
            logger.info(f"Starting transcription for {self.file.name} ({self.file.size} bytes)")
            
            # Create a temporary file with the correct extension
            suffix = f'.{file_ext}'
            
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                try:
                    # Save the file content to the temporary file
                    logger.info("Writing file to temporary location")
                    for chunk in self.file.chunks():
                        tmp.write(chunk)
                    tmp.flush()
                    
                    # Check file size - limit transcription for very large files
                    file_size_mb = os.path.getsize(tmp.name) / (1024 * 1024)
                    if file_size_mb > 100:  # Limit to 100MB files
                        logger.warning(f"File too large for transcription: {file_size_mb:.1f}MB")
                        return f"File too large for transcription ({file_size_mb:.1f}MB). Maximum size is 100MB."
                    
                    # Load Whisper model - use smaller model for efficiency
                    try:
                        if not whisper:
                            return "Audio transcription requires OpenAI Whisper. Install with: pip install openai-whisper"
                        
                        logger.info("Loading Whisper model")
                        
                        # Use 'tiny' model for faster processing, 'base' for better accuracy
                        model_size = 'tiny' if file_size_mb > 10 else 'base'
                        model = whisper.load_model(model_size)
                        
                        logger.info(f"Starting transcription with {model_size} model")
                        result = model.transcribe(
                            tmp.name,
                            verbose=False,  # Reduce console output
                            word_timestamps=False,  # Faster processing
                            fp16=False  # Better compatibility
                        )
                        
                        transcript = result.get("text", "").strip()
                        logger.info(f"Transcription completed: {len(transcript)} characters")
                        
                        return transcript
                        
                    except Exception as transcription_error:
                        error_msg = f"Transcription failed: {str(transcription_error)}"
                        logger.error(error_msg)
                        return error_msg
                    
                finally:
                    # Clean up the temporary file
                    try:
                        os.unlink(tmp.name)
                        logger.info("Temporary file cleaned up")
                    except Exception as cleanup_error:
                        logger.warning(f"Failed to clean up temporary file: {cleanup_error}")
                        
        except Exception as e:
            logger.error(f"Media transcription failed for {self.file.name}: {str(e)}")
            return f"Transcription failed: {str(e)}"

    def extract_text(self):
        """Extract text content from supported files with improved error handling"""
        import logging
        logger = logging.getLogger(__name__)
        
        if not self.file:
            return ""
            
        try:
            file_ext = self.file.name.lower().split('.')[-1] if self.file.name else ''
            logger.info(f"Extracting text from {file_ext} file")
            
            if file_ext == 'pdf':
                return self._extract_pdf_text()
            elif file_ext in ['docx', 'doc']:
                return self._extract_docx_text()
            elif file_ext in ['txt', 'md', 'html', 'htm']:
                return self._extract_plain_text()
            else:
                logger.info(f"No text extraction method for .{file_ext} files")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {self.file.name}: {e}")
            return f"Text extraction failed: {str(e)}"

    def _extract_plain_text(self):
        """Extract text from plain text files"""
        try:
            content = ""
            with self.file.open('r', encoding='utf-8') as f:
                content = f.read()
            return content
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with self.file.open('r', encoding='latin-1') as f:
                    content = f.read()
                return content
            except Exception:
                return "Could not decode file content"
        except Exception as e:
            return f"Error reading plain text file: {str(e)}"

    def _extract_pdf_text(self):
        """Extract text from PDF files with better error handling"""
        try:
            import PyPDF2
            import tempfile
            
            with tempfile.NamedTemporaryFile(delete=True) as tmp:
                # Write file content to temp file
                for chunk in self.file.chunks():
                    tmp.write(chunk)
                tmp.flush()
                tmp.seek(0)
                
                # Extract text
                reader = PyPDF2.PdfReader(tmp)
                text_content = []
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        text_content.append(f"[Error extracting page {page_num + 1}: {str(e)}]")
                
                return '\n\n'.join(text_content)
                
        except ImportError:
            return "PDF text extraction requires PyPDF2 library. Please install it with: pip install PyPDF2"
        except Exception as e:
            return f"PDF extraction failed: {str(e)}"

    def _extract_docx_text(self):
        """Extract text from DOCX files with better error handling"""
        try:
            from docx import Document as DocxDocument
            import tempfile
            
            with tempfile.NamedTemporaryFile(delete=True, suffix='.docx') as tmp:
                # Write file content to temp file
                for chunk in self.file.chunks():
                    tmp.write(chunk)
                tmp.flush()
                tmp.seek(0)
                
                # Extract text
                doc = DocxDocument(tmp)
                paragraphs = []
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            paragraphs.append(' | '.join(row_text))
                
                return '\n\n'.join(paragraphs)
                
        except ImportError:
            return "DOCX text extraction requires python-docx library. Please install it with: pip install python-docx"
        except Exception as e:
            return f"DOCX extraction failed: {str(e)}"
    
    def extract_html_content(self):
        """Extract HTML content from DOCX files for editing"""
        if not self.file.name.lower().endswith(('.docx', '.doc')):
            return ""
        
        try:
            from documents.docx_converter import convert_docx_to_html
            
            # Read file content
            file_content = b''
            for chunk in self.file.chunks():
                file_content += chunk
            
            return convert_docx_to_html(file_content)
            
        except Exception as e:
            print(f"Error extracting HTML from DOCX: {e}")
            return ""
    
    def update_from_html(self, html_content: str) -> bool:
        """
        Update document content from HTML and create a new DOCX file
        
        Args:
            html_content: Updated HTML content
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            from documents.docx_converter import convert_html_to_docx
            from django.core.files.base import ContentFile
            import os
            
            # Convert HTML back to DOCX
            docx_content = convert_html_to_docx(html_content, self.title)
            
            if not docx_content:
                print("Failed to convert HTML to DOCX - no content generated")
                return False
            
            # Get original filename, handle both forward and backward slashes
            if self.file and self.file.name:
                filename = os.path.basename(self.file.name)
            else:
                filename = f"{self.title}.docx"
            
            # Ensure filename has correct extension
            if not filename.lower().endswith(('.docx', '.doc')):
                # Remove existing extension and add .docx
                name_without_ext = os.path.splitext(filename)[0]
                filename = f"{name_without_ext}.docx"
            
            # Create ContentFile with the DOCX content
            content_file = ContentFile(docx_content)
            
            # Save new DOCX content (this will replace the existing file)
            self.file.save(
                filename,
                content_file,
                save=False  # Don't save the model instance yet
            )
            
            # Update HTML content and extracted text
            self.html_content = html_content
            # Re-extract text from the new DOCX file
            self.content = self.extract_text()
            
            return True
            
        except Exception as e:
            import traceback
            print(f"Error updating document from HTML: {e}")
            print(f"Traceback: {traceback.format_exc()}")
            return False

    def convert_to_pdf(self, save_file=True) -> bool:
        """
        Convert the Word document (.docx) to PDF format
        
        Args:
            save_file: Whether to save the PDF file to the model field
            
        Returns:
            bool: True if conversion was successful, False otherwise
        """
        import os
        import tempfile
        from django.core.files.base import ContentFile
        
        # Check if we have a .docx file to convert
        if not self.file or not self.file.name.lower().endswith(('.docx', '.doc')):
            print(f"Cannot convert to PDF: file is not a Word document")
            return False
        
        try:
            # Create temporary directories
            temp_dir = tempfile.mkdtemp()
            input_path = os.path.join(temp_dir, f"input_{self.pk}.docx")
            output_path = os.path.join(temp_dir, f"output_{self.pk}.pdf")
            
            # Copy the original file to temp directory
            with open(input_path, 'wb') as temp_file:
                self.file.open('rb')
                temp_file.write(self.file.read())
                self.file.close()
            
            # Try different conversion methods
            conversion_success = False
            
            # Method 1: Try docx2pdf (Windows only)
            try:
                from docx2pdf import convert
                convert(input_path, output_path)
                conversion_success = True
                print("PDF conversion successful using docx2pdf")
            except ImportError:
                print("docx2pdf not available, trying alternative methods...")
            except Exception as e:
                print(f"docx2pdf conversion failed: {e}")
            
            # Method 2: Try pypandoc (cross-platform, requires pandoc installation)
            if not conversion_success:
                try:
                    import pypandoc
                    pypandoc.convert_file(input_path, 'pdf', outputfile=output_path)
                    conversion_success = True
                    print("PDF conversion successful using pypandoc")
                except ImportError:
                    print("pypandoc not available")
                except Exception as e:
                    print(f"pypandoc conversion failed: {e}")
            
            # Method 3: Try python-docx + reportlab (fallback method)
            if not conversion_success:
                try:
                    conversion_success = self._convert_via_reportlab(input_path, output_path)
                    if conversion_success:
                        print("PDF conversion successful using reportlab fallback")
                except Exception as e:
                    print(f"Reportlab conversion failed: {e}")
            
            if conversion_success and os.path.exists(output_path):
                if save_file:
                    # Read the generated PDF and save to model
                    with open(output_path, 'rb') as pdf_file:
                        pdf_content = pdf_file.read()
                    
                    # Generate PDF filename
                    original_name = os.path.splitext(os.path.basename(self.file.name))[0]
                    pdf_filename = f"{original_name}.pdf"
                    
                    # Save PDF to model field
                    self.pdf_file.save(
                        pdf_filename,
                        ContentFile(pdf_content),
                        save=False  # Don't trigger save() recursion
                    )
                
                # Clean up temp files
                try:
                    os.unlink(input_path)
                    os.unlink(output_path)
                    os.rmdir(temp_dir)
                except:
                    pass
                
                return True
            
            else:
                print("PDF conversion failed: output file not created")
                return False
                
        except Exception as e:
            print(f"Error during PDF conversion: {e}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
            return False
    
    def _convert_via_reportlab(self, input_path: str, output_path: str) -> bool:
        """
        Fallback PDF conversion using python-docx + reportlab
        """
        try:
            from docx import Document as DocxDocument
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            
            # Read the Word document
            doc = DocxDocument(input_path)
            
            # Create PDF
            pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Extract text from Word document and add to PDF
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    story.append(Paragraph(paragraph.text, styles['Normal']))
                    story.append(Spacer(1, 0.2*inch))
            
            # Build PDF
            pdf_doc.build(story)
            return True
            
        except ImportError:
            print("ReportLab not available for fallback conversion")
            return False
        except Exception as e:
            print(f"ReportLab conversion error: {e}")
            return False


class DocumentVersion(models.Model):
    document = models.ForeignKey(
        Document,
        related_name='versions',
        on_delete=models.CASCADE,
        verbose_name=_('Document')
    )
    file = models.FileField(
        upload_to=document_upload_path,
        verbose_name=_('File')
    )
    version_number = models.PositiveIntegerField(verbose_name=_('Version Number'))
    changes = models.TextField(blank=True, verbose_name=_('Change Description'))
    created_by = models.ForeignKey(
        User,
        on_delete=models.PROTECT,
        verbose_name=_('Created By')
    )
    created_at = models.DateTimeField(default=now, verbose_name=_('Created At'))
    approved = models.BooleanField(default=False, verbose_name=_('Approved'))

    class Meta:
        ordering = ['-version_number']
        unique_together = ('document', 'version_number')
        verbose_name = _('Document Version')
        verbose_name_plural = _('Document Versions')

    def __str__(self):
        return f"v{self.version_number} of {self.document.title}"

    def clean(self):
        if self.version_number < 1:
            raise ValidationError(_("Version number must be at least 1"))


class Tag(models.Model):
    name = models.CharField(max_length=50, unique=True, verbose_name=_('Name'))
    created_by = models.ForeignKey(
        User,
        on_delete=models.PROTECT,
        verbose_name=_('Created By'))
    created_at = models.DateTimeField(default=now, verbose_name=_('Created At'))
    is_system = models.BooleanField(default=False, verbose_name=_('System Tag'))

    class Meta:
        ordering = ['name']
        verbose_name = _('Tag')
        verbose_name_plural = _('Tags')

    def __str__(self):
        return self.name


class AuditLog(models.Model):
    ACTION_CHOICES = [
        ('CREATE', 'Created'),
        ('UPDATE', 'Updated'),
        ('VIEW', 'Viewed'),
        ('DOWNLOAD', 'Downloaded'),
        ('APPROVE', 'Approved'),
        ('REJECT', 'Rejected'),
        ('SUBMIT_REVIEW', 'Submitted for Review'),
        ('VERSION_CREATE', 'New Version Created'),
        ('CONTENT_EDIT', 'Content Edited'),
        ('STATUS_CHANGE', 'Status Changed'),
        ('ARCHIVE', 'Archived'),
        ('RESTORE', 'Restored'),
        ('DELETE', 'Permanently Deleted'),
        ('SHARE', 'Shared'),
        ('PERMISSION_CHANGE', 'Permissions Changed'),
    ]

    user = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        verbose_name=_('User')
    )
    # Document reference fields
    document_id = models.PositiveIntegerField(null=True)
    document_title = models.CharField(max_length=255, blank=True)
    document_type = models.CharField(max_length=20, blank=True)
    
    action = models.CharField(max_length=20, choices=ACTION_CHOICES)
    timestamp = models.DateTimeField(auto_now_add=True)
    additional_info = models.JSONField(default=dict)
    
    # Additional context fields
    ip_address = models.GenericIPAddressField(null=True, blank=True)
    user_agent = models.TextField(blank=True)
    session_key = models.CharField(max_length=40, blank=True)

    class Meta:
        indexes = [
            models.Index(fields=['document_id']),
            models.Index(fields=['user', 'timestamp']),
            models.Index(fields=['action', 'timestamp']),
        ]
        ordering = ['-timestamp']
        verbose_name = _('Audit Log Entry')
        verbose_name_plural = _('Audit Log Entries')

    def __str__(self):
        user_name = self.user.get_full_name() if self.user else 'Anonymous'
        return f"{user_name} {self.get_action_display().lower()} document {self.document_id} at {self.timestamp.strftime('%Y-%m-%d %H:%M')}"
    
    @classmethod
    def log_action(cls, user, document, action, request=None, **additional_info):
        """
        Convenience method to create audit log entries
        """
        log_data = {
            'user': user,
            'document_id': document.pk if document else None,
            'document_title': document.title if document else '',
            'document_type': document.get_document_type_display() if document else '',
            'action': action,
            'additional_info': additional_info,
        }
        
        if request:
            log_data.update({
                'ip_address': cls._get_client_ip(request),
                'user_agent': request.META.get('HTTP_USER_AGENT', ''),
                'session_key': request.session.session_key or '',
            })
        
        return cls.objects.create(**log_data)
    
    @staticmethod
    def _get_client_ip(request):
        """Get client IP address from request"""
        x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
        if x_forwarded_for:
            ip = x_forwarded_for.split(',')[0]
        else:
            ip = request.META.get('REMOTE_ADDR')
        return ip
    
from django.db import models
from ckeditor.fields import RichTextField

class DocumentTemplate(models.Model):
    name = models.CharField(max_length=100)
    content = RichTextField()
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    def __str__(self):
        return self.name